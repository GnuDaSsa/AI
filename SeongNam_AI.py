import streamlit as st
import re
from openai import OpenAI
from pymongo import MongoClient
from datetime import datetime
from docx import Document
from odf.text import P, Span, H, ListItem
from odf.opendocument import load
import io


st.set_page_config(layout="wide",page_title="성남시 인공지능 자동화 서비스")

client = OpenAI() #####API키를 세팅하세요######

#MongoDB Cloudtype 클라이언트 접속정보 ####### 데이터 베이스 서버에 연결하세요 #####
#client_db = MongoClient('mongodb://') #Cloudtype에 등록된 서비스끼지 접속방법
client_db = MongoClient('mongodb://')  #외부에서 접속하는 방법

#MongoDB 보도자료 데이터베이스
press_db = client_db['SeongnamPress']  # 성남시 보도자료 데이터베이스
press_release_collection = press_db['Seongnam_releases']  # 보도자료를 저장할 컬렉션

#MongoDB 카운터 데이터베이스
count_db = client_db['SeongNamCounter']
counter_collection = count_db['SeongNamPpressCounter']
stats_collection = count_db['SeongNamPressStats']  # 총 카운트를 저장할 별도 컬렉션

def 보도자료_생성(제목, 내용, 소감, 담당자, 연락처): 
    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"당신은 배태랑 기자입니다. 입력된 {제목}을 바탕으로 보도자료 제목 5개를 번호를 붙혀서 추천하고, 이를 '보도자료 추천 제목'이라는 진한글씨로 제목 아래 각각 ''로 묶어서 화면에 보여주세요. 각 추천 제목은 한 줄씩 출력해야합니다. 그 후 '{제목}'을 주제로 보도자료를 작성하세요. 보도자료 내용은 '{내용}'을 기반으로 주제를 정하여 2500자 이상 풍부하게 작성합니다. 보도자료의 마지막 단락에는 '{소감}의 내용을 출력하고 소감을 풍부하게 작성하고, 담당자와 연락처는 각각 '{담당자}','{연락처}'로 작성합니다. 각 단락은 간결하고 공식적으로 작성해야 하며, 전형적인 보도자료 구조로 문장의 종결은 '이다','된다','했다'로 마무리합니다. "}, 
        ],
        stream=True
    )
    return completion

def 보도자료_출력(completion):
    output_placeholder = st.empty()
    full_text = ""
    for chunk in completion:
        if hasattr(chunk.choices[0].delta, 'content'):
            content = chunk.choices[0].delta.content
            if content:
                full_text += content
                output_placeholder.write(full_text)
    return full_text

def 보도자료_저장(제목, 내용, 담당부서, 담당자, 연락처, ai_content):
    press_release_collection.insert_one({
        "title": 제목,
        "summary": 내용,
        "department": 담당부서,
        "person_in_charge": 담당자,
        "contact": 연락처,
        "ai_generated_content": ai_content
    })

def 최신_보도자료_불러오기():
    latest_press_release = press_release_collection.find_one({}, sort=[("_id", -1)])  # 최신 문서 불러오기
    return latest_press_release

def 보도자료_중복_확인(제목, 내용):
    # 주어진 제목과 내용이 기존 보도자료와 중복되는지 확인
    중복_보도자료 = press_release_collection.find_one({"title": 제목, "summary": 내용})
    return 중복_보도자료 is not None

# 카운터 초기화 함수
def 카운터_초기화():
    today_str = datetime.now().strftime('%Y-%m-%d')  # 문자열 형식으로 날짜 변환
    counter_collection.update_one(
        {"date": today_str},
        {"$setOnInsert": {"count": 0}},  # 오늘의 카운트 초기화
        upsert=True
    )
    # 총 카운트 초기화
    stats_collection.update_one(
        {"type": "total_count"},
        {"$setOnInsert": {"value": 0}},
        upsert=True
    )

# 카운터 증가 함수
def 카운터_증가():
    today_str = datetime.now().strftime('%Y-%m-%d')  # 문자열 형식으로 날짜 변환

    # 오늘의 카운트를 증가
    counter_collection.find_one_and_update(
        {"date": today_str},
        {"$inc": {"count": 1}},  # 오늘의 카운트 증가
        return_document=True
    )
    # 총 카운트를 증가
    stats_collection.find_one_and_update(
        {"type": "total_count"},
        {"$inc": {"value": 1}},  # 총 카운트 증가
        return_document=True
    )
    # 업데이트된 총 카운트 반환
    total_count_entry = stats_collection.find_one({"type": "total_count"})
    return total_count_entry['value']


def replace_text_in_node(node, search_text, replace_text):
    """노드를 재귀적으로 탐색하여 텍스트를 변경하는 함수"""
    if node.nodeType == node.TEXT_NODE:
        if search_text in node.data:
            node.data = node.data.replace(search_text, str(replace_text)) 
    else:
        for child in node.childNodes:
            replace_text_in_node(child, search_text, replace_text)

def replace_text_in_elements(doc, search_text, replace_text):
    """텍스트 문서에서 모든 단락, 제목, 스팬 및 리스트 아이템 내 텍스트를 변경하는 함수"""
    elements = [P, Span, H, ListItem]
    for elem in elements:
        for node in doc.getElementsByType(elem):
            replace_text_in_node(node, search_text, replace_text)

def add_paragraph(doc, text):
    """ODT 문서에 단락을 추가하는 함수"""
    paragraph = P(text=text)
    doc.text.addElement(paragraph)


def 오늘_생성한_보도자료_불러오기():
    # 최근 10개의 인사말씀을 내림차순으로 가져오기
    최신_press_releases = list(press_release_collection.find().sort("_id", -1).limit(count_entry['count']))
    return 최신_press_releases


# 카운터 초기화
카운터_초기화()

# 당일 카운트 조회
today_str = datetime.now().strftime('%Y-%m-%d')  # 문자열 형식으로 날짜 변환
count_entry = counter_collection.find_one({"date": today_str}) # 오늘 카운트 조회
total_count_entry = stats_collection.find_one({"type": "total_count"}) # 총 누적 카운트 조회

# 현재 날짜 및 시간 얻기
now = datetime.now()

# 요일을 한국어로 매핑
days_kor = ['월', '화', '수', '목', '금', '토', '일']
weekday_kor = days_kor[now.weekday()]

# 원하는 형식으로 변환 (요일을 추가)
formatted_date = now.strftime('%Y.%m.%d') + f'.({weekday_kor})'


st.title("성남시 생성형 AI 보도자료 작성 서비스(ChatGPT-4o)")

# 사이드바에서 최근 10개의 보도자료을 선택
with st.sidebar:
    st.subheader("**📢 성남시 보도자료 작성 서비스**")
    최신_press_releases = 오늘_생성한_보도자료_불러오기()
    titles = [release['title'] for release in 최신_press_releases]

    if titles:
        selected_title = st.selectbox("**🌈 오늘 생성된 보도자료를 확인해보세요**", titles)
    else:
        st.write("생성된 보도자료가 없습니다.")

col1, col2 = st.columns([8.5,1.5]) 
with col1:
    st.warning("**1. 단어와 간단한 문장을 입력하는 것만으로도 여러분의 아이디어와 정보를 효과적으로 전달할 수 있습니다.**\n\n" 
               "**2. 복잡한 과정이 필요 없이, 직관적인 사용 방식으로 신속하게 보도자료를 완성해 보세요!**\n\n")

with col2: 
    st.error(f"**✔️오늘 {count_entry['count']}회**\n\n **✔️누적 {total_count_entry['value']}회**")

col3, col4 = st.columns([5.5,4.5])

# 선택한 보도자료 데이터를 가져오기
selected_release = next((release for release in 최신_press_releases if release['title'] == selected_title), None)

with col3:
    text_container = st.empty()
    with text_container.container():
        # 최신 보도자료 불러오기
        if selected_release:
            st.subheader(selected_release['title'])
            st.write(selected_release['ai_generated_content'])
        else:
            st.write("생성된 보도자료가 없습니다.")        

with col4:
    
    제목col, 부서명col = st.columns([6,4])
    with 제목col:
        제목 = st.text_input("**1.보도자료 제목**")
        if 제목:
            if re.search(r"[^\w\s'\".]", 제목):  # 단어 문자(\w), 공백(\s), 작은따옴표('), 큰따옴표("), 점(.) 외의 문자 검색
                st.warning("특수 문자가 포함되어 있습니다. %와 같은 문자는 사용할 수 없습니다.")
    with 부서명col:
        담당부서 = st.text_input("**2.담당부서**", placeholder='4차산업추진국 AI반도체과')
            
    col5, col6, col7 = st.columns(3)
    with col5:        
        소감 = st.text_input("**3.소감주체**", placeholder='성남시장 OOO')
    with col6:        
        담당자 = st.text_input("**4.주무관**", placeholder='홍길동')
    with col7:            
        연락처 = st.text_input("**5.연락처**", placeholder='031-729-0000')

    내용 = st.text_area("**6.보도자료 핵심반영 내용**", height=170,placeholder='(예시)물놀이장 이용은 초등학생 이하로 성남시민은 신분증 확인후 우선 입장하고, 운영시간은 오전 10시부터 오후 5시까지며 매주 월요일은 휴무입니다. 다양한 물놀이 시설과 안전관리자 간호사 배치')


    col17, col18 = st.columns([2.5,7.5])    
    with col17:
    #보도자료 생성 버튼 클릭 이벤트 핸들링
        if st.button("**:red[보도자료 생성]**"):
            if 제목 and 담당부서 and 내용 and 담당자 and 연락처:
                if 보도자료_중복_확인(제목, 내용):
                    with col4:  
                        st.error("동일한 제목과 내용으로 생성된 보도자료가 이미 존재합니다.\n\n" "다른 내용으로 시도해 주세요.")
                else:
                    with col3:
                        text_container.empty()
                        completion = 보도자료_생성(제목, 내용, 소감, 담당자, 연락처)
                        st.subheader(제목)
                        full_text = 보도자료_출력(completion)
                        보도자료_저장(제목, 내용, 담당부서, 담당자, 연락처, full_text)  # 생성된 보도자료 저장
                        카운터_증가()
                        st.rerun()
            else:
                st.error("모든 필드를 입력해주세요.")

    if selected_release is not None:
        if 'summary' in selected_release:
            summary = f"**사용자가 입력한 프롬프트를 참고하세요.**\n\n {selected_release['summary']}"
        else:
            summary = "선택된 보도자료에 요약 정보가 없습니다."
        st.info(summary, icon=None)
    else:
        st.info("선택된 보도자료가 없습니다.", icon=None)


    with col18:
        # 선택된 보도자료기준 문서 생성
        if selected_release:
            title = selected_release['title'].replace(" ", "_")[:15]
            Press_path = "SeongNam_Press.odt"  # 수정할 파일 경로

            try:
                # ODT 파일 생성
                with open(Press_path, "rb") as file:
                    odt_file = file.read()
                doc = load(io.BytesIO(odt_file))
                replace_text_in_elements(doc, "입력부서", selected_release['department'])
                replace_text_in_elements(doc, "입력담당자", selected_release['person_in_charge'])
                replace_text_in_elements(doc, "입력연락처", selected_release['contact'])
                replace_text_in_elements(doc, "입력제목", selected_release['title'])
                replace_text_in_elements(doc, "입력일자", formatted_date)

                
                add_paragraph(doc, " ") # 공백진입

                입력본문 = selected_release['ai_generated_content']
                paragraphs = 입력본문.strip().split('\n')
                for i, paragraph in enumerate(paragraphs):
                    add_paragraph(doc, paragraph)

                odt_out_file = io.BytesIO()
                doc.save(odt_out_file)
                odt_out_file.seek(0)
                # 다운로드 버튼 추가
                st.download_button(
                    label=f"{title}보도자료 다운로드",
                    data=odt_out_file,
                    file_name=f"{title}.odt",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"문서 파일 생성 중 오류 발생: {e}")
        else:
            st.error("선택된 보도자료가 없습니다.")
